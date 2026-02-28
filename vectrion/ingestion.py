"""Vectorian Data Ingestion Engine

Processes uploaded evidence files for breach response engagements.
Supports structured, semi-structured, unstructured, image (OCR), and video (frame OCR) inputs.
All processing is local — no data leaves the server.
"""
from __future__ import annotations

import hashlib
import json
import re
from pathlib import Path
from typing import Any

# ─────────────────────────────────────────────────────────────────────────────
# EXTENSION SETS
# ─────────────────────────────────────────────────────────────────────────────
_STRUCTURED = {".csv", ".tsv", ".json", ".jsonl", ".xml", ".xlsx", ".xls"}
_DOCUMENT   = {".pdf", ".docx", ".doc", ".txt", ".md", ".rtf"}
_WEB        = {".html", ".htm", ".eml"}
_IMAGE      = {".jpg", ".jpeg", ".png", ".gif", ".bmp", ".tiff", ".tif", ".webp"}
_VIDEO      = {".mp4", ".mov", ".avi", ".mkv", ".wmv", ".flv", ".webm"}

ALLOWED_EXTENSIONS = _STRUCTURED | _DOCUMENT | _WEB | _IMAGE | _VIDEO

TYPE_ICONS = {
    "structured": "📊",
    "document":   "📄",
    "web":        "🌐",
    "image":      "🖼",
    "video":      "🎥",
    "unknown":    "📁",
}

# ─────────────────────────────────────────────────────────────────────────────
# UTILITIES
# ─────────────────────────────────────────────────────────────────────────────

def file_hash(path: Path) -> str:
    """SHA-256 of file for chain-of-custody logging."""
    h = hashlib.sha256()
    with open(path, "rb") as f:
        for chunk in iter(lambda: f.read(65536), b""):
            h.update(chunk)
    return h.hexdigest()


def detect_type(path: Path) -> tuple[str, str]:
    """Returns (category, human-readable label)."""
    ext = Path(path).suffix.lower()
    if ext in _STRUCTURED:  return "structured", "Structured Data"
    if ext in _DOCUMENT:    return "document",   "Document"
    if ext in _WEB:         return "web",        "Web / Email"
    if ext in _IMAGE:       return "image",      "Image"
    if ext in _VIDEO:       return "video",      "Video"
    return "unknown", "Unknown"


# ─────────────────────────────────────────────────────────────────────────────
# STRUCTURED DATA EXTRACTORS
# ─────────────────────────────────────────────────────────────────────────────

def _extract_csv(path: Path) -> tuple[str, str]:
    import csv
    rows: list[str] = []
    with open(path, newline="", encoding="utf-8", errors="replace") as f:
        reader = csv.reader(f)
        for i, row in enumerate(reader):
            if i >= 2000:
                rows.append(f"[... {i} more rows truncated]")
                break
            rows.append("\t".join(str(v) for v in row))
    return "\n".join(rows), "csv_parse"


def _extract_json(path: Path) -> tuple[str, str]:
    with open(path, encoding="utf-8", errors="replace") as f:
        content = f.read()
    if path.suffix.lower() == ".jsonl":
        lines = [ln for ln in content.splitlines() if ln.strip()][:1000]
        return "\n".join(lines), "jsonl_parse"
    try:
        obj = json.loads(content)
        return json.dumps(obj, indent=2)[:100_000], "json_parse"
    except Exception:
        return content[:100_000], "json_raw"


def _extract_xml(path: Path) -> tuple[str, str]:
    try:
        import xml.etree.ElementTree as ET
        tree = ET.parse(path)
        root = tree.getroot()
        parts: list[str] = []

        def collect(el: Any, depth: int = 0) -> None:
            tag = el.tag.split("}")[-1] if "}" in el.tag else el.tag
            text = (el.text or "").strip()
            parts.append("  " * depth + f"<{tag}>" + (f" {text}" if text else ""))
            for child in el:
                collect(child, depth + 1)

        collect(root)
        return "\n".join(parts[:10_000]), "xml_parse"
    except Exception:
        with open(path, encoding="utf-8", errors="replace") as f:
            return f.read()[:100_000], "xml_raw"


def _extract_xlsx(path: Path) -> tuple[str, str]:
    try:
        import openpyxl
        wb = openpyxl.load_workbook(path, data_only=True)
        rows: list[str] = []
        for sheet_name in wb.sheetnames[:10]:
            ws = wb[sheet_name]
            rows.append(f"=== Sheet: {sheet_name} ===")
            for i, row in enumerate(ws.iter_rows(values_only=True)):
                if i >= 2000:
                    rows.append("[... more rows truncated]")
                    break
                rows.append("\t".join("" if v is None else str(v) for v in row))
        return "\n".join(rows), "xlsx_openpyxl"
    except ImportError:
        return "[XLSX extraction unavailable — install openpyxl]", "unavailable"
    except Exception as e:
        return f"[XLSX extraction error: {e}]", "error"


# ─────────────────────────────────────────────────────────────────────────────
# DOCUMENT EXTRACTORS
# ─────────────────────────────────────────────────────────────────────────────

def _extract_pdf(path: Path) -> tuple[str, str]:
    # 1. pdfplumber (best for structured PDFs with tables)
    try:
        import pdfplumber
        pages: list[str] = []
        with pdfplumber.open(path) as pdf:
            for i, page in enumerate(pdf.pages[:200]):
                text = page.extract_text() or ""
                if text.strip():
                    pages.append(f"--- Page {i + 1} ---\n{text}")
        text = "\n\n".join(pages)
        if text.strip():
            return text, "pdfplumber"
    except ImportError:
        pass
    except Exception:
        pass

    # 2. pypdf fallback
    try:
        from pypdf import PdfReader
        reader = PdfReader(path)
        pages = []
        for i, page in enumerate(reader.pages[:200]):
            t = page.extract_text() or ""
            if t.strip():
                pages.append(f"--- Page {i + 1} ---\n{t}")
        text = "\n\n".join(pages)
        if text.strip():
            return text, "pypdf"
    except ImportError:
        pass
    except Exception:
        pass

    # 3. OCR fallback
    return _ocr_pdf_fallback(path)


def _ocr_pdf_fallback(path: Path) -> tuple[str, str]:
    try:
        from pdf2image import convert_from_path
        import pytesseract
        images = convert_from_path(str(path), dpi=150, first_page=1, last_page=20)
        pages = [pytesseract.image_to_string(img) for img in images]
        return "\n\n".join(pages), "pdf_ocr"
    except ImportError:
        return "[PDF text extraction unavailable — install pdfplumber or pytesseract+pdf2image]", "unavailable"
    except Exception as e:
        return f"[PDF OCR error: {e}]", "error"


def _extract_docx(path: Path) -> tuple[str, str]:
    try:
        from docx import Document
        doc = Document(path)
        parts: list[str] = []
        for para in doc.paragraphs:
            if para.text.strip():
                parts.append(para.text)
        for table in doc.tables:
            for row in table.rows:
                parts.append("\t".join(c.text for c in row.cells))
        return "\n".join(parts), "python_docx"
    except ImportError:
        return "[DOCX extraction unavailable — install python-docx]", "unavailable"
    except Exception as e:
        return f"[DOCX extraction error: {e}]", "error"


def _extract_html(path: Path) -> tuple[str, str]:
    with open(path, encoding="utf-8", errors="replace") as f:
        raw = f.read()
    try:
        from bs4 import BeautifulSoup
        soup = BeautifulSoup(raw, "html.parser")
        for tag in soup(["script", "style", "head"]):
            tag.decompose()
        return soup.get_text(separator="\n"), "beautifulsoup"
    except ImportError:
        text = re.sub(r"<[^>]+>", " ", raw)
        text = re.sub(r"\s+", " ", text).strip()
        return text, "html_regex"
    except Exception as e:
        return f"[HTML extraction error: {e}]", "error"


def _extract_text(path: Path) -> tuple[str, str]:
    with open(path, encoding="utf-8", errors="replace") as f:
        return f.read(), "plaintext"


# ─────────────────────────────────────────────────────────────────────────────
# IMAGE OCR
# ─────────────────────────────────────────────────────────────────────────────

def _extract_image(path: Path) -> tuple[str, str]:
    try:
        import pytesseract
        from PIL import Image
        img = Image.open(path)
        if img.mode not in ("RGB", "L", "RGBA"):
            img = img.convert("RGB")
        text = pytesseract.image_to_string(img)
        return text, "tesseract_ocr"
    except ImportError:
        return "[Image OCR unavailable — install Pillow and pytesseract]", "unavailable"
    except Exception as e:
        return f"[Image OCR error: {e}]", "error"


# ─────────────────────────────────────────────────────────────────────────────
# VIDEO FRAME OCR
# ─────────────────────────────────────────────────────────────────────────────

def _extract_video(path: Path) -> tuple[str, str]:
    """Sample one frame every 5 seconds, OCR each, return concatenated text."""

    # Try opencv-python first
    try:
        import cv2
        import pytesseract
        from PIL import Image

        cap = cv2.VideoCapture(str(path))
        fps = cap.get(cv2.CAP_PROP_FPS) or 25
        sample_every = max(1, int(fps * 5))
        texts: list[str] = []
        i = 0
        while cap.isOpened():
            ret, frame = cap.read()
            if not ret:
                break
            if i % sample_every == 0:
                img = Image.fromarray(cv2.cvtColor(frame, cv2.COLOR_BGR2RGB))
                t = pytesseract.image_to_string(img).strip()
                if t:
                    texts.append(f"[Frame {i}]\n{t}")
                if len(texts) >= 60:
                    texts.append("[...video truncated after 60 text frames...]")
                    break
            i += 1
            if i > fps * 600:  # cap at 10 minutes
                break
        cap.release()
        return "\n\n".join(texts) or "[No readable text found in video frames]", "video_ocr_cv2"

    except ImportError:
        pass
    except Exception as e:
        return f"[Video OCR error (cv2): {e}]", "error"

    # Try imageio[ffmpeg] as alternative
    try:
        import imageio
        import pytesseract
        from PIL import Image

        reader = imageio.get_reader(str(path), "ffmpeg")
        meta = reader.get_meta_data()
        fps = meta.get("fps", 25)
        sample_every = max(1, int(fps * 5))
        texts: list[str] = []
        for i, frame in enumerate(reader):
            if i % sample_every == 0:
                img = Image.fromarray(frame)
                if img.mode != "RGB":
                    img = img.convert("RGB")
                t = pytesseract.image_to_string(img).strip()
                if t:
                    texts.append(f"[Frame {i}]\n{t}")
                if len(texts) >= 60:
                    break
            if i > fps * 600:
                break
        reader.close()
        return "\n\n".join(texts) or "[No readable text found in video frames]", "video_ocr_imageio"

    except ImportError:
        return (
            "[Video OCR unavailable — install opencv-python or imageio[ffmpeg] with pytesseract]",
            "unavailable",
        )
    except Exception as e:
        return f"[Video OCR error (imageio): {e}]", "error"


# ─────────────────────────────────────────────────────────────────────────────
# PUBLIC API
# ─────────────────────────────────────────────────────────────────────────────

def process_file(path: Path) -> dict[str, Any]:
    """
    Process a file and extract its text content.

    Returns a dict with keys:
        type        - category (structured/document/web/image/video/unknown)
        type_label  - human-readable category
        text        - extracted text (may be empty on error)
        method      - extraction method used
        sha256      - SHA-256 hash of original file (chain of custody)
        size_bytes  - file size in bytes
        error       - error message or None
    """
    path = Path(path)
    cat, label = detect_type(path)
    sha = file_hash(path)
    size = path.stat().st_size
    ext = path.suffix.lower()

    try:
        if ext in (".csv", ".tsv"):
            text, method = _extract_csv(path)
        elif ext in (".json", ".jsonl"):
            text, method = _extract_json(path)
        elif ext == ".xml":
            text, method = _extract_xml(path)
        elif ext in (".xlsx", ".xls"):
            text, method = _extract_xlsx(path)
        elif ext == ".pdf":
            text, method = _extract_pdf(path)
        elif ext in (".docx", ".doc"):
            text, method = _extract_docx(path)
        elif ext in (".html", ".htm", ".eml"):
            text, method = _extract_html(path)
        elif ext in (".txt", ".md", ".rtf"):
            text, method = _extract_text(path)
        elif ext in _IMAGE:
            text, method = _extract_image(path)
        elif ext in _VIDEO:
            text, method = _extract_video(path)
        else:
            text, method = "[Unsupported file type]", "none"

        return {
            "type": cat,
            "type_label": label,
            "text": text,
            "method": method,
            "sha256": sha,
            "size_bytes": size,
            "error": None,
        }

    except Exception as e:
        return {
            "type": cat,
            "type_label": label,
            "text": "",
            "method": "error",
            "sha256": sha,
            "size_bytes": size,
            "error": str(e),
        }
